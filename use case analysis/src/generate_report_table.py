import os
import re
import copy
import pandas as pd
from collections import Counter
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def clean_filename(text):
    """
    Clean and format the section heading text to be used as a filename.
    E.g. " 01   Contexte de la mission" -> "[REPORT] AI builders - 01 - Contexte de la mission.docx"
    """
    match = re.match(r'\s*(\d+)\s+(.+)', text)
    if match:
        num = match.group(1)
        name = match.group(2).strip()
        name = re.sub(r'[\\/*?:"<>|]', '_', name)
        name = re.sub(r'\s+', ' ', name)
        return f"[REPORT] AI builders - {num} - {name}.docx"
    else:
        cleaned = re.sub(r'[\\/*?:"<>|]', '_', text.strip())
        cleaned = re.sub(r'\s+', ' ', cleaned)
        return f"[REPORT] AI builders - {cleaned}.docx"

def add_element_to_doc(doc, element):
    """
    Copy a block-level XML element into the body of another document,
    ensuring it is inserted before the final sectPr element.
    """
    new_element = copy.deepcopy(element)
    body = doc.element.body
    sectPr = body.xpath('w:sectPr')
    if sectPr:
        sectPr[0].addprevious(new_element)
    else:
        body.append(new_element)

def main():
    # Setup paths relative to the script directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    excel_path = os.path.join(script_dir, "..", "output", "use_cases_catalog.xlsx")
    docx_path = os.path.join(script_dir, "..", "output", "[REPORT] AI builders.docx")
    output_dir = os.path.join(script_dir, "..", "output")

    print("======================================================================")
    print("STEP 1: GENERATING DYNAMIC DATA SOURCE TABLE")
    print("======================================================================")
    
    print(f"Reading Excel catalog from: {excel_path}")
    if not os.path.exists(excel_path):
        print(f"Error: Excel file not found at {excel_path}")
        return

    # Load and process data
    df = pd.read_excel(excel_path, sheet_name="Catalogue")
    
    # Filter for Medium and Large complexity tiers
    medium_df = df[df['Complexity_Tier'] == 'Medium']
    large_df = df[df['Complexity_Tier'] == 'Large']
    combined_df = df[df['Complexity_Tier'].isin(['Medium', 'Large'])]

    def count_sources(dataframe):
        counts = Counter()
        for _, row in dataframe.iterrows():
            sources_str = row['Data_Sources']
            if pd.isna(sources_str):
                continue
            
            # Split by comma and strip whitespaces
            sources = [s.strip() for s in str(sources_str).split(',')]
            for source in sources:
                source_clean = source.strip()
                # Exclude empty values and "A revoir avec le builder"
                if source_clean and source_clean.lower() != "a revoir avec le builder":
                    counts[source_clean] += 1
        return counts

    medium_counts = count_sources(medium_df)
    large_counts = count_sources(large_df)
    combined_counts = count_sources(combined_df)

    # Sort sources by descending total count
    sorted_sources = sorted(list(combined_counts.keys()), key=lambda x: combined_counts[x], reverse=True)

    print(f"Found {len(sorted_sources)} unique data sources (excluding 'A revoir avec le builder')")

    print(f"Opening main Word document: {docx_path}")
    if not os.path.exists(docx_path):
        print(f"Error: Word document not found at {docx_path}")
        return

    doc = docx.Document(docx_path)

    # Idempotency check: look for our existing intro paragraph and delete it along with the table
    intro_marker = "Répartition des sources de données unitaires pour les cas d'usage complexes (Medium & Large)"
    
    found_idx = -1
    for idx, para in enumerate(doc.paragraphs):
        if intro_marker in para.text:
            found_idx = idx
            break

    if found_idx != -1:
        print("Found existing table section. Removing it first for clean regeneration...")
        p_to_remove = doc.paragraphs[found_idx]
        p_element = p_to_remove._element
        
        # Check if the next element is a table
        next_elm = p_element.getnext()
        if next_elm is not None and next_elm.tag.endswith('tbl'):
            next_elm.getparent().remove(next_elm)
            
        p_element.getparent().remove(p_element)
        print("Existing section removed.")

    # Find the target paragraph to insert after
    # We want to insert after the paragraph containing "03   Packager les sources de données (Accélérateur IT)"
    # specifically after the long description that follows it
    target_header_text = "03   Packager les sources de données (Accélérateur IT)"
    header_idx = -1
    for idx, para in enumerate(doc.paragraphs):
        if target_header_text in para.text:
            header_idx = idx
            break

    if header_idx == -1:
        print(f"Warning: Could not find paragraph with header: '{target_header_text}'")
        # Fallback to the end of the document if header not found
        insert_after_para = doc.paragraphs[-1]
    else:
        # The paragraph immediately following the header is the long description (Para 174)
        insert_after_para = doc.paragraphs[header_idx + 1]
        print(f"Found target paragraph: '{insert_after_para.text[:50]}...'")

    # Helper function to insert a paragraph after another paragraph
    def insert_paragraph_after(para, text, style=None):
        new_p = para._parent.add_paragraph(text, style=style)
        para._element.addnext(new_p._element)
        return new_p

    # Helper function to insert a table after a paragraph
    def insert_table_after(para, rows, cols):
        table = para._parent.add_table(rows, cols, Inches(6.5))
        p_element = para._element
        t_element = table._element
        p_element.addnext(t_element)
        return table

    # Helper function to set cell shading (background color)
    def set_cell_background(cell, fill_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    # Helper function to set cell margins (padding)
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        cell._tc.get_or_add_tcPr().append(tcMar)

    # Helper function to set elegant horizontal-only borders
    def set_table_borders(table):
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            borders = tblPr[0].xpath('w:tblBorders')
            for border in borders:
                tblPr[0].remove(border)
            new_borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D9D9D9"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D9D9D9"/>'
                f'  <w:left w:val="none"/>'
                f'  <w:right w:val="none"/>'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
                f'  <w:insideV w:val="none"/>'
                f'</w:tblBorders>'
            )
            tblPr[0].append(new_borders)

    # 1. Insert the introduction paragraph
    intro_p = insert_paragraph_after(
        insert_after_para, 
        f"Répartition des sources de données unitaires pour les cas d'usage complexes (Medium & Large) :"
    )
    intro_p.style = 'Normal'
    intro_p.paragraph_format.space_before = Pt(18)
    intro_p.paragraph_format.space_after = Pt(8)
    intro_p.paragraph_format.keep_with_next = True
    
    # Format the intro paragraph run
    run = intro_p.runs[0] if intro_p.runs else intro_p.add_run()
    run.font.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    # Use dark blue color matching GAS style: #1B4F72 (RGB: 27, 79, 114)
    run.font.color.rgb = docx.shared.RGBColor(27, 79, 114)

    # 2. Insert the table (1 header row + N data rows, 4 columns)
    num_rows = len(sorted_sources) + 1
    table = insert_table_after(intro_p, num_rows, 4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Dual width definition (US Letter content width is 6.5 inches = 9360 DXA)
    # Column widths: Source (3.2 inches), Medium (1.1 inches), Large (1.1 inches), Total (1.1 inches)
    col_widths = [Inches(3.2), Inches(1.1), Inches(1.1), Inches(1.1)]

    # Style Header Row
    headers = ["Source de données unitaire", "Medium", "Large", "Total"]
    hdr_cells = table.rows[0].cells
    for col_idx, text in enumerate(headers):
        cell = hdr_cells[col_idx]
        cell.width = col_widths[col_idx]
        set_cell_background(cell, "1B4F72") # Deep blue header #1B4F72
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150) # Padding
        
        # Paragraph styling inside cell
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = docx.shared.RGBColor(255, 255, 255) # White text

    # Style Data Rows
    for row_idx, source in enumerate(sorted_sources):
        row_cells = table.rows[row_idx + 1].cells
        
        m_val = medium_counts.get(source, 0)
        l_val = large_counts.get(source, 0)
        t_val = combined_counts.get(source, 0)
        
        row_data = [source, str(m_val), str(l_val), str(t_val)]
        
        # Alternating background colors: White and very light gray-blue (#F8FAFC)
        bg_hex = "FFFFFF" if row_idx % 2 == 0 else "F8FAFC"
        
        for col_idx, text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150) # Padding
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            
            # Make the source name or the counts look extra professional
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = docx.shared.RGBColor(60, 60, 60)
            else:
                run.font.color.rgb = docx.shared.RGBColor(80, 80, 80)
                if col_idx == 3: # Total column
                    run.font.bold = True
                    run.font.color.rgb = docx.shared.RGBColor(27, 79, 114) # Highlight total in dark blue

    # Save document
    print(f"Saving modified document to: {docx_path}")
    doc.save(docx_path)
    print("Main document saved successfully with dynamic data source table!")

    print("\n======================================================================")
    print("STEP 2: SPLITTING REPORT INTO 11 SEGMENTED DOCUMENTS")
    print("======================================================================")

    # Re-open the saved document to have the exact updated content for splitting
    doc = docx.Document(docx_path)

    sections = []
    current_section = {
        "title": "00_Intro",
        "elements": []
    }
    
    # Loop sequentially through all children of the body
    for element in doc.element.body:
        if element.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(element, doc)
            if p.style.name == "Heading 1":
                # Save previous section if it has elements
                if current_section["elements"]:
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    "title": p.text,
                    "elements": [element]
                }
                continue
            
            # Skip empty section breaks or standard section properties
            if element.tag.endswith('sectPr'):
                continue
                
            current_section["elements"].append(element)
            
        elif element.tag.endswith('tbl'):
            current_section["elements"].append(element)
            
    # Add final section
    if current_section["elements"]:
        sections.append(current_section)

    # Merge cover page / intro paragraphs (before the first Heading 1) into Section 1
    if sections and sections[0]["title"] == "00_Intro":
        intro_section = sections.pop(0)
        if sections:
            sections[0]["elements"] = intro_section["elements"] + sections[0]["elements"]

    print(f"Dividing into {len(sections)} sections...")

    # Process and write each of the 11 section documents
    for idx, sec in enumerate(sections):
        filename = clean_filename(sec["title"])
        dest_path = os.path.join(output_dir, filename)
        
        # Load fresh template to preserve styles, margins, headers, footers
        sec_doc = docx.Document(docx_path)
        
        # Clear body of the template
        for p in list(sec_doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(sec_doc.tables):
            t._element.getparent().remove(t._element)
            
        # Add elements for this specific section
        for element in sec["elements"]:
            add_element_to_doc(sec_doc, element)
            
        # Save split document
        sec_doc.save(dest_path)
        print(f" -> Created Section {idx+1:02d}: {filename}")

    print("\nAll 11 section documents generated successfully in the output directory!")
    print("======================================================================")

if __name__ == "__main__":
    main()
