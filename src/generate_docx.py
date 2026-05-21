#!/usr/bin/env python3
"""
generate_docx.py
────────────────
Génère le rapport de synthèse Word (.docx) Air Liquide — AI Champions
avec la charte graphique Pyl.Tech (navy #0b132b, jaune #F4BF46, turquoise #208AAE).

Style : Rapport de consulting — cabinet Tier-1 (méthodologie, cadre d'analyse,
        recommandations structurées, narrative professionnelle).

Usage : python3 src/generate_docx.py
Output: output/rapport_ai_champions.docx
"""
import sys
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import pandas as pd
except ImportError as e:  # pragma: no cover
    print(f"❌ Dépendance manquante : {e}\n   pip install python-docx pandas openpyxl")  # noqa
    sys.exit(1)

# ── Chemins ───────────────────────────────────────────────────────────────────
PROJECT_DIR = Path(__file__).parent.parent
CATALOG_FILE = PROJECT_DIR / "output" / "use_cases_catalog.xlsx"
OUTPUT_FILE = PROJECT_DIR / "output" / "rapport_ai_champions.docx"

# ── Palette Pyl.Tech officielle ───────────────────────────────────────────────
PYL_NAVY_DARK = RGBColor(0x0B, 0x13, 0x2B)
PYL_NAVY = RGBColor(0x0D, 0x21, 0x49)
PYL_YELLOW = RGBColor(0xF4, 0xBF, 0x46)
PYL_TEAL_BLUE = RGBColor(0x20, 0x8A, 0xAE)
PYL_TEAL = RGBColor(0x5B, 0xC0, 0xBE)
PYL_BODY_GREY = RGBColor(0x4F, 0x4F, 0x4F)
PYL_GREY_BG = RGBColor(0xEE, 0xEE, 0xEE)
PYL_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PYL_SUCCESS = RGBColor(0x13, 0x86, 0x36)
PYL_DANGER = RGBColor(0xC9, 0x14, 0x32)

YEAR = date.today().year

# ── Helpers XML ───────────────────────────────────────────────────────────────


def set_cell_bg(cell, rgb: RGBColor) -> None:
    """Applique un fond coloré à une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_val = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:fill"), hex_val)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_border_left(paragraph, hex_color: str, size_pt: int = 18) -> None:
    """Ajoute une bordure gauche colorée sur un paragraphe (callout)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def highlight_run(run, hex_color: str = "F4BF46") -> None:
    """Surlignage jaune charte (style marqueur fluo)."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    rPr.append(shd)


def set_paragraph_spacing(
    paragraph, space_before: int = 0, space_after: int = 6
) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), str(space_before * 20))
    pSp.set(qn("w:after"), str(space_after * 20))
    pPr.append(pSp)


# ── Composants de document ────────────────────────────────────────────────────


def apply_base_styles(doc: Document) -> None:
    """Configure les styles globaux du document."""
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Poppins"
    normal.font.size = Pt(12)
    normal.font.color.rgb = PYL_BODY_GREY

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Poppins"
    h1.font.bold = True
    h1.font.size = Pt(24)
    h1.font.color.rgb = PYL_NAVY_DARK

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Poppins"
    h2.font.bold = True
    h2.font.size = Pt(18)
    h2.font.color.rgb = PYL_NAVY_DARK

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Poppins"
    h3.font.bold = True
    h3.font.size = Pt(14)
    h3.font.color.rgb = PYL_NAVY_DARK


def add_cover_page(doc: Document) -> None:
    """Page de couverture style Pyl.Tech consulting."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, PYL_YELLOW)
    cell.width = Cm(16)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Air Liquide — AI Champions")
    run.font.name = "Poppins"
    run.font.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = PYL_NAVY_DARK
    set_paragraph_spacing(p, 24, 4)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Analyse, Classification et Recommandations\ndu Portefeuille Use Cases IA")
    r2.font.name = "Poppins"
    r2.font.italic = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = PYL_NAVY_DARK
    set_paragraph_spacing(p2, 4, 24)

    # Sous-bloc info
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Produit par Pyl.Tech  |  {date.today().strftime('%B %Y')}")
    r3.font.name = "Poppins"
    r3.font.size = Pt(11)
    r3.font.color.rgb = PYL_TEAL_BLUE
    set_paragraph_spacing(p3, 12, 4)

    # Mention confidentiel
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Document confidentiel — Usage interne Air Liquide")
    r4.font.name = "Poppins"
    r4.font.size = Pt(10)
    r4.font.color.rgb = PYL_BODY_GREY
    r4.font.italic = True
    set_paragraph_spacing(p4, 4, 12)

    doc.add_page_break()


def add_section_header(
    doc: Document, number: str, title: str, subtitle: str = ""
) -> None:
    """Titre de section avec cartouche jaune."""
    p = doc.add_paragraph()
    p.style = "Heading 1"
    set_paragraph_spacing(p, 24, 6)

    badge = p.add_run(f" {number} ")
    badge.font.name = "Poppins"
    badge.font.bold = True
    badge.font.size = Pt(14)
    badge.font.color.rgb = PYL_NAVY_DARK
    highlight_run(badge)

    p.add_run("  ")

    title_run = p.add_run(title)
    title_run.font.name = "Poppins"
    title_run.font.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = PYL_NAVY_DARK

    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.style = "Normal"
        p2.runs[0].font.size = Pt(12)
        p2.runs[0].font.color.rgb = PYL_BODY_GREY
        set_paragraph_spacing(p2, 0, 8)


def add_kpi_block(doc: Document, kpis: list[tuple[str, str, str]]) -> None:
    """Bloc KPIs en tableau style Pyl.Tech."""
    tbl = doc.add_table(rows=1 + len(kpis), cols=3)
    tbl.style = "Table Grid"

    headers = ["Indicateur", "Valeur", "Contexte"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, PYL_NAVY_DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Poppins"
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = PYL_WHITE

    for r_idx, (label, value, context) in enumerate(kpis):
        bg = PYL_NAVY if r_idx % 2 == 0 else PYL_WHITE
        fg = PYL_WHITE if r_idx % 2 == 0 else PYL_BODY_GREY

        for c_idx, text in enumerate([label, value, context]):
            cell = tbl.cell(r_idx + 1, c_idx)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if c_idx == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = p.add_run(text)
            run.font.name = "Poppins"
            run.font.size = Pt(10)
            run.font.color.rgb = fg
            if c_idx == 1:
                run.font.bold = True

    doc.add_paragraph()


def add_pivot_table(
    doc: Document, title: str, headers: list[str], rows: list[list[str]]
) -> None:
    """Tableau croisé Pyl.Tech."""
    if title:
        p_title = doc.add_paragraph(title)
        p_title.style = "Heading 3"
        set_paragraph_spacing(p_title, 12, 4)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, PYL_NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Poppins"
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = PYL_WHITE

    for r_idx, row_data in enumerate(rows):
        bg = PYL_GREY_BG if r_idx % 2 == 0 else PYL_WHITE
        fg = PYL_BODY_GREY
        for c_idx, val in enumerate(row_data):
            cell = tbl.cell(r_idx + 1, c_idx)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if c_idx == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = p.add_run(str(val))
            run.font.name = "Poppins"
            run.font.size = Pt(10)
            run.font.color.rgb = fg
            if c_idx == len(row_data) - 1:
                run.font.bold = True

    doc.add_paragraph()


def add_callout(doc: Document, text: str, style: str = "info") -> None:
    """Encadré callout avec bordure gauche colorée."""
    color_map = {
        "info": "F4BF46",
        "warning": "C91432",
        "success": "138636",
        "note": "0b132b",
    }
    border_hex = color_map.get(style, "F4BF46")

    p = doc.add_paragraph()
    add_border_left(p, border_hex, size_pt=24)
    run = p.add_run(text)
    run.font.name = "Poppins"
    run.font.size = Pt(11)
    run.font.color.rgb = PYL_BODY_GREY
    run.font.italic = True
    set_paragraph_spacing(p, 6, 6)


def add_bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    """Bullet Pyl.Tech avec puce carrée jaune."""
    p = doc.add_paragraph()
    p.style = "Normal"

    bullet_run = p.add_run("  ")
    bullet_run.font.name = "Poppins"
    bullet_run.font.size = Pt(12)
    bullet_run.font.color.rgb = PYL_YELLOW
    bullet_run.font.bold = True

    if bold_prefix:
        bold_run = p.add_run(bold_prefix)
        bold_run.font.name = "Poppins"
        bold_run.font.bold = True
        bold_run.font.size = Pt(12)
        bold_run.font.color.rgb = PYL_NAVY_DARK

        sep = p.add_run(" — ")
        sep.font.name = "Poppins"
        sep.font.size = Pt(12)
        sep.font.color.rgb = PYL_BODY_GREY

    body_run = p.add_run(text)
    body_run.font.name = "Poppins"
    body_run.font.size = Pt(12)
    body_run.font.color.rgb = PYL_BODY_GREY
    set_paragraph_spacing(p, 0, 4)


def add_body_text(doc: Document, text: str) -> None:
    """Paragraphe de corps de texte standard."""
    p = doc.add_paragraph(text)
    p.style = "Normal"
    set_paragraph_spacing(p, 4, 8)


def add_sub_heading(doc: Document, text: str) -> None:
    """Sous-titre Heading 2."""
    p = doc.add_paragraph(text)
    p.style = "Heading 2"
    set_paragraph_spacing(p, 16, 6)


# ── Corps du document ─────────────────────────────────────────────────────────


def build_document(doc: Document, df: pd.DataFrame) -> None:
    """Construit le contenu complet du rapport consulting."""

    # ── Métriques globales ────────────────────────────────────────────────────
    total = len(df)
    tiers = df["Complexity_Tier"].value_counts().to_dict()
    it_count = (df["IT_Flag"] != "").sum()
    clusters = df["Cluster"].nunique()
    job_families = df["Job Family"].nunique()
    nb_tools = df["Nb_Tools"].mean()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 01 : CONTEXTE DE LA MISSION
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "01",
        "Contexte de la mission",
        "Cadrage du projet et objectifs de l'analyse",
    )

    add_body_text(
        doc,
        "Dans le cadre de son programme AI Champions, Air Liquide a identifié et mobilisé "
        "des profils métiers — non-IT — pour explorer et déployer des solutions d'intelligence "
        "artificielle dans leurs activités quotidiennes. Ces champions, répartis sur l'ensemble "
        "des clusters géographiques et des familles métier du groupe, ont collectivement produit "
        f"un portefeuille de {total} use cases IA."
    )

    add_body_text(
        doc,
        "Pyl.Tech a été mandaté pour réaliser une analyse structurée de ce portefeuille, "
        "avec quatre objectifs :"
    )

    add_bullet(
        doc,
        "Regrouper les use cases en familles fonctionnelles transverses, "
        "au-delà des silos organisationnels.",
        "Comprendre le corpus",
    )
    add_bullet(
        doc,
        "Scorer chaque use case selon 5 dimensions de complexité "
        "(intégration, périmètre, données, IA, impact) pour établir une classification "
        "Small / Medium / Large.",
        "Qualifier la complexité",
    )
    add_bullet(
        doc,
        "Proposer des architectures de référence Google-first adaptées à chaque "
        "niveau de complexité, en identifiant les points de dépendance IT.",
        "Projeter la cible",
    )
    add_bullet(
        doc,
        "Formuler des recommandations concrètes pour structurer, documenter et "
        "pérenniser les réalisations des champions.",
        "Guider la pratique",
    )

    add_callout(
        doc,
        "Ce rapport constitue le livrable principal de la mission. Les références aux use cases "
        "utilisent uniquement des identifiants UC_ID anonymisés (aucun nom de personne).",
        style="note",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 02 : MÉTHODOLOGIE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "02",
        "Methodologie d'analyse",
        "Approche en 6 phases pour qualifier le portefeuille AI Champions",
    )

    add_body_text(
        doc,
        "L'analyse a suivi une méthodologie structurée en 6 phases séquentielles, "
        "conçue pour traiter un corpus hétérogène de use cases soumis par des profils non-IT, "
        "avec des niveaux de maturité et de documentation variables."
    )

    # Phase 1 — Audit
    add_sub_heading(doc, "Phase 1 — Audit et nettoyage des donnees")
    add_body_text(
        doc,
        "Le fichier source (Advanced AI Champions - Action Monitoring.xlsx) a fait l'objet "
        "d'un audit systématique. Les anomalies identifiées — valeurs #REF! (7 lignes), "
        "stages non renseignés (141 lignes), impacts économiques absents (193 lignes) — "
        "ont été traitées par des règles conservatrices (remplacement par des valeurs neutres, "
        "aucune exclusion). Chaque use case a reçu un identifiant unique UC_ID."
    )

    add_pivot_table(
        doc,
        "Anomalies identifiees et traitement",
        ["Anomalie", "Volume", "Decision"],
        [
            ["Valeurs #REF! (Cluster, Job Family)", "7 lignes", "Remplace par N/A"],
            ["Stage non renseigne", "141 lignes", "Valeur : Non renseigne"],
            ["Impact economique absent", "193 lignes", "Valeur : Non evalue"],
            ["Outils multi-valeurs", "Toutes lignes", "Normalisation en tags"],
            ["Descriptions identiques, clusters differents", "Variable", "UC_ID partage"],
        ],
    )

    # Phase 2 — Classification fonctionnelle
    add_sub_heading(doc, "Phase 2 — Classification en familles fonctionnelles")
    add_body_text(
        doc,
        "Une analyse sémantique des descriptions longues, croisée avec les Job Families "
        "et les outils déclarés, a permis d'identifier 7 familles fonctionnelles transverses. "
        "Cette taxonomie dépasse les silos organisationnels pour révéler les grands patterns "
        "d'usage de l'IA dans le groupe."
    )

    add_pivot_table(
        doc,
        "Les 7 familles fonctionnelles",
        ["Code", "Famille", "Description"],
        [
            ["F1", "Automatisation documentaire", "Generation, traduction, resume, redaction de documents"],
            ["F2", "Assistant BI et decisionnel", "Dashboards, analyses predictives, alertes KPI"],
            ["F3", "Customer & Sales Intelligence", "Analyse clients, optimisation visites, scoring CRM"],
            ["F4", "Monitoring & Maintenance", "Prediction de pannes, analyse capteurs, surveillance process"],
            ["F5", "Knowledge Management", "FAQ, onboarding, bases de connaissances, chatbots formation"],
            ["F6", "Automatisation de workflows", "Scripts, flows, integrations Google/Office"],
            ["F7", "Data Engineering & Reporting", "Pipelines de donnees, ETL, visualisation, Power BI"],
        ],
    )

    # Phase 3 — Scoring
    add_sub_heading(doc, "Phase 3 — Scoring de complexite multi-dimensionnel")
    add_body_text(
        doc,
        "Chaque use case a été scoré sur 5 dimensions indépendantes, chacune notée de 1 à 3 points "
        "(score total : 5 à 15). Ce scoring multi-dimensionnel permet une classification objective "
        "et reproductible, indépendante du jugement subjectif."
    )

    add_pivot_table(
        doc,
        "Grille de scoring (5 dimensions x 3 niveaux)",
        ["Dimension", "1 pt (Faible)", "2 pts (Moyen)", "3 pts (Eleve)"],
        [
            [
                "D1 — Integration technique",
                "1 outil, no-code (L1/L2)",
                "2-3 outils, semi-code (L3)",
                "4+ outils, code custom (L4)",
            ],
            [
                "D2 — Perimetre organisationnel",
                "Equipe locale",
                "Country / Cluster",
                "Group (deploiement global)",
            ],
            [
                "D3 — Complexite data",
                "Donnees statiques / manuelles",
                "Donnees connectees (SFDC, BI)",
                "Donnees temps reel / industrielles",
            ],
            [
                "D4 — Maturite IA requise",
                "Prompting / Gem / NotebookLM",
                "API Gemini + RAG basique",
                "Agent multi-etapes / ML",
            ],
            [
                "D5 — Impact economique declare",
                "Non evalue / Productivite",
                "Cost Reduction",
                "Revenue Growth / Sustainability",
            ],
        ],
    )

    add_body_text(
        doc,
        "Le score total determine le tier de complexite :"
    )

    add_pivot_table(
        doc,
        "Classification par tier",
        ["Tier", "Score", "Label", "Time-to-value", "Profil"],
        [
            ["Small", "5-7", "Quick Win", "< 2 semaines", "Champion seul, no-code"],
            ["Medium", "8-11", "Use Case Structurant", "4-8 semaines", "Champion + IT local"],
            ["Large", "12-15", "Projet Strategique", "3-12 mois", "Equipe IT + Champion"],
        ],
    )

    add_callout(
        doc,
        "Regle DONE-marker : lorsqu'une description contient un marqueur de maturite explicite "
        "(\"DONE SO FAR\", \"Already built\"), le scoring prospectif porte uniquement sur le scope "
        "futur restant a realiser. Cela evite de sur-scorer des use cases dont la partie complexe "
        "est deja implementee.",
        style="info",
    )

    # Phase 4 — Signalement IT
    add_sub_heading(doc, "Phase 4 — Detection des dependances IT")
    add_body_text(
        doc,
        "Un algorithme de détection automatique identifie les use cases impliquant des systèmes "
        "d'entreprise (SFDC, SAP, DCS, SCADA, Active Directory, OAuth, API keys, Cloud Run, "
        "Vertex AI, BigQuery). Ces use cases reçoivent un drapeau IT obligatoire, signalant "
        "qu'ils ne peuvent pas être conduits par un champion seul."
    )

    # Phase 5 — Architecture cible
    add_sub_heading(doc, "Phase 5 — Architecture cible Google-first")
    add_body_text(
        doc,
        "Pour chaque tier de complexité, une architecture de référence a été définie, "
        "en restant dans l'écosystème Google Workspace / GCP. Tout besoin dépassant cet "
        "écosystème est signalé comme point d'attention IT."
    )

    # Phase 6 — Recommandations
    add_sub_heading(doc, "Phase 6 — Recommandations et bonnes pratiques")
    add_body_text(
        doc,
        "Enfin, des recommandations concrètes ont été formulées pour structurer, documenter "
        "et pérenniser les réalisations des champions. Ces recommandations sont adaptées au profil "
        "non-IT des champions et visent à maximiser la réutilisabilité et la maintenabilité "
        "des use cases produits."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 03 : EXECUTIVE SUMMARY (SYNTHESE)
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "03",
        "Synthese du portefeuille",
        "Vue d'ensemble des resultats de l'analyse",
    )

    add_body_text(
        doc,
        f"L'analyse du portefeuille AI Champions porte sur {total} use cases valides, "
        f"repartis sur {clusters} clusters et {job_families} familles metier. "
        f"La majorite des initiatives ({tiers.get('Small', 0) + tiers.get('Medium', 0)} use cases, "
        f"soit {(tiers.get('Small', 0) + tiers.get('Medium', 0)) * 100 // total}%) "
        f"presente un profil accessible (Small ou Medium), confirmant une maturite collective "
        f"significative dans l'adoption des outils IA Google Workspace."
    )

    add_body_text(
        doc,
        f"{it_count} use cases ({it_count * 100 // total}%) necessitent un accompagnement IT "
        f"dedie en raison de dependances systemes (SFDC, SAP, DCS, SCADA, API d'entreprise). "
        f"Seuls {tiers.get('Large', 0)} projets ({tiers.get('Large', 0) * 100 // total}%) "
        f"sont classes Large, confirmant que les champions IA restent dans un perimetre "
        f"d'autonomie maitrise."
    )

    kpis = [
        ("Use cases analyses", str(total), "Portefeuille complet"),
        ("Clusters couverts", str(clusters), "Perimetre geographique/organisationnel"),
        ("Familles metier", str(job_families), "Diversite des profils champions"),
        (
            "Quick Wins (Small)",
            str(tiers.get("Small", 0)),
            f"{tiers.get('Small', 0) * 100 // total}% — Deployables en < 2 semaines",
        ),
        (
            "Use cases structurants (Medium)",
            str(tiers.get("Medium", 0)),
            f"{tiers.get('Medium', 0) * 100 // total}% — 4 a 8 semaines de delivery",
        ),
        (
            "Projets strategiques (Large)",
            str(tiers.get("Large", 0)),
            f"{tiers.get('Large', 0) * 100 // total}% — 3 a 12 mois, IT implique",
        ),
        ("Points d'attention IT", str(it_count), "Escalade IT obligatoire"),
        ("Nb moyen d'outils par UC", f"{nb_tools:.1f}", "Proxy de complexite d'integration"),
    ]
    add_kpi_block(doc, kpis)

    add_callout(
        doc,
        "La population de champions AI n'est pas une population IT. Les use cases Large ou "
        "impliquant des systemes d'entreprise (SFDC, SAP, DCS, SCADA) ne peuvent etre conduits "
        "sans accompagnement technique dedie. Ce constat structure l'ensemble de nos recommandations.",
        style="warning",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 04 : REPARTITION PAR COMPLEXITE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "04",
        "Repartition par complexite",
        "Classification Small / Medium / Large selon 5 dimensions de scoring",
    )

    add_body_text(
        doc,
        "Le scoring multi-dimensionnel (5 dimensions, score 5-15) segmente le portefeuille "
        "en trois tiers de complexite, chacun avec un profil d'accompagnement et un time-to-value "
        "distincts."
    )

    add_pivot_table(
        doc,
        "Repartition globale par tier",
        ["Tier", "Score", "Label", "Time-to-value", "Nb Use Cases", "Part"],
        [
            [
                "Small", "5-7", "Quick Win", "< 2 semaines",
                str(tiers.get("Small", 0)),
                f"{tiers.get('Small', 0) * 100 // total}%",
            ],
            [
                "Medium", "8-11", "Use Case Structurant", "4-8 semaines",
                str(tiers.get("Medium", 0)),
                f"{tiers.get('Medium', 0) * 100 // total}%",
            ],
            [
                "Large", "12-15", "Projet Strategique", "3-12 mois",
                str(tiers.get("Large", 0)),
                f"{tiers.get('Large', 0) * 100 // total}%",
            ],
            ["TOTAL", "", "", "", str(total), "100%"],
        ],
    )

    # Famille x Tier
    add_sub_heading(doc, "Croisement Famille fonctionnelle x Tier")
    pivot_rows = []
    for fam_label in sorted(df["Family_Label"].unique()):
        subset = df[df["Family_Label"] == fam_label]
        s = subset["Complexity_Tier"].value_counts().to_dict()
        tot = len(subset)
        pivot_rows.append(
            [
                fam_label,
                str(s.get("Small", 0)),
                str(s.get("Medium", 0)),
                str(s.get("Large", 0)),
                str(tot),
            ]
        )
    pivot_rows.sort(key=lambda x: -int(x[4]))
    pivot_rows.append(
        [
            "TOTAL",
            str(tiers.get("Small", 0)),
            str(tiers.get("Medium", 0)),
            str(tiers.get("Large", 0)),
            str(total),
        ]
    )

    add_pivot_table(
        doc,
        "Famille x Complexite",
        ["Famille fonctionnelle", "Small", "Medium", "Large", "Total"],
        pivot_rows,
    )

    add_body_text(
        doc,
        "Observation cle : les familles F3 (Customer & Sales Intelligence) et F7 "
        "(Data Engineering & Reporting) concentrent la majorite des use cases Medium et Large, "
        "ce qui reflete la complexite intrinseque des integrations CRM et des pipelines de donnees."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 05 : ANALYSE PAR CLUSTER
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "05",
        "Analyse par Cluster",
        "Repartition geographique et organisationnelle des use cases",
    )

    add_body_text(
        doc,
        "La distribution par cluster révèle des niveaux de maturité et d'engagement variables "
        "selon les entités. Certains clusters (Airgas, NCE, GBU InnoTech) présentent une densité "
        "élevée de use cases, reflétant une adoption avancée du programme AI Champions."
    )

    cluster_data = (
        df.groupby("Cluster")["Complexity_Tier"]
        .value_counts()
        .unstack(fill_value=0)
    )
    cluster_rows = []
    for cluster in cluster_data.index:
        row = [cluster]
        for tier in ["Small", "Medium", "Large"]:
            val = (
                int(cluster_data.loc[cluster, tier])
                if tier in cluster_data.columns
                else 0
            )
            row.append(str(val))
        row.append(str(int(cluster_data.loc[cluster].sum())))
        cluster_rows.append(row)
    cluster_rows.sort(key=lambda x: -int(x[4]))

    add_pivot_table(
        doc,
        "Cluster x Complexite (Top 15)",
        ["Cluster", "Small", "Medium", "Large", "Total"],
        cluster_rows[:15],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 06 : CATALOGUE PAR FAMILLE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "06",
        "Catalogue par famille fonctionnelle",
        "Presentation detaillee des 7 familles — definition, volumetrie et Quick Wins",
    )

    family_descriptions = {
        "Automatisation documentaire": (
            "F1",
            "Cette famille regroupe les use cases lies a la generation, "
            "la traduction, le resume et la redaction automatisee de documents. "
            "Les outils dominants sont Gemini (Prompts/Gems) et NotebookLM, "
            "positionnes sur du no-code/low-code accessibles aux champions.",
        ),
        "Assistant BI & decisionnel": (
            "F2",
            "Ces use cases exploitent l'IA pour augmenter les capacites d'analyse "
            "et de decision : tableaux de bord intelligents, analyses predictives, "
            "alertes automatiques sur KPI. Power BI, App Script et Gemini sont les "
            "outils les plus utilises.",
        ),
        "Customer & Sales Intelligence": (
            "F3",
            "Famille dense et complexe couvrant l'analyse clients SFDC, "
            "l'optimisation des tournees commerciales, le scoring client et le CRM augmente. "
            "C'est la famille avec le plus grand nombre de use cases Large, "
            "en raison des integrations systemes requises.",
        ),
        "Monitoring & Maintenance industrielle": (
            "F4",
            "Use cases lies a la surveillance des equipements industriels, "
            "la prediction de pannes, l'analyse de capteurs (DCS, SCADA, AVEVA). "
            "Cette famille concentre les cas les plus techniquement complexes "
            "et les plus dependants de l'IT.",
        ),
        "Knowledge Management & Formation": (
            "F5",
            "Bases de connaissances intelligentes, FAQ automatisees, chatbots de formation, "
            "onboarding assiste par IA. NotebookLM est l'outil de reference pour ces use cases, "
            "offrant un acces no-code aux champions.",
        ),
        "Automatisation de workflows internes": (
            "F6",
            "Automatisation de processus administratifs et operationnels via "
            "App Script, Workspace Studio (ex-Flows) et AppSheet. "
            "Cette famille presente le meilleur ratio Small/Medium, "
            "confirmant son accessibilite aux champions non-IT.",
        ),
        "Data Engineering & Reporting": (
            "F7",
            "Famille la plus technique : pipelines de donnees, extraction/transformation, "
            "reporting avance via Power BI et Python (Fabric, DataStudio). "
            "Les competences requises depassent souvent le profil champion standard, "
            "necessitant un track 'Data Champion' dedie.",
        ),
    }

    for fam_label, (fam_code, fam_desc) in family_descriptions.items():
        subset = df[df["Family_Label"] == fam_label]
        if subset.empty:
            continue

        p_fam = doc.add_paragraph()
        p_fam.style = "Heading 2"
        badge = p_fam.add_run(f" {fam_code} ")
        badge.font.name = "Poppins"
        badge.font.bold = True
        badge.font.size = Pt(12)
        badge.font.color.rgb = PYL_NAVY_DARK
        highlight_run(badge)
        p_fam.add_run("  ")
        tr = p_fam.add_run(fam_label)
        tr.font.name = "Poppins"
        tr.font.bold = True
        tr.font.size = Pt(16)
        tr.font.color.rgb = PYL_NAVY_DARK
        set_paragraph_spacing(p_fam, 16, 4)

        add_body_text(doc, fam_desc)

        tier_s = subset["Complexity_Tier"].value_counts().to_dict()
        add_bullet(doc, f"{len(subset)} use cases identifies", "Volume")
        add_bullet(
            doc,
            f"{tier_s.get('Small', 0)} Small / {tier_s.get('Medium', 0)} Medium / {tier_s.get('Large', 0)} Large",
            "Repartition",
        )
        it_sub = (subset["IT_Flag"] != "").sum()
        if it_sub > 0:
            add_bullet(
                doc,
                f"{it_sub} use cases avec point d'attention IT",
                "Alerte IT",
            )

        # Top Quick Wins
        quick_wins_uc = subset[subset["Complexity_Tier"] == "Small"].head(3)
        if not quick_wins_uc.empty:
            p_h = doc.add_paragraph("Top Quick Wins de la famille")
            p_h.style = "Heading 3"
            set_paragraph_spacing(p_h, 8, 4)
            for _, row in quick_wins_uc.iterrows():
                desc_short = str(row["Use Case Description (Long)"])[:120].strip()
                if len(str(row["Use Case Description (Long)"])) > 120:
                    desc_short += "..."
                add_bullet(
                    doc,
                    f"[{row['UC_ID']}] {desc_short} (Score: {row['Score_Total']})",
                )

        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 07 : ARCHITECTURE CIBLE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "07",
        "Architectures de reference",
        "Stack cible Google-first par niveau de complexite",
    )

    add_body_text(
        doc,
        "Pour chaque tier de complexité, une architecture de référence a été définie "
        "en privilégiant l'écosystème Google Workspace et GCP. L'objectif est de fournir "
        "aux champions un cadre technique clair, avec une identification explicite des "
        "moments où l'intervention IT devient nécessaire."
    )

    # Small
    add_sub_heading(doc, "Small — Prompting & Automation")
    add_body_text(
        doc,
        "Les use cases Small reposent entierement sur les outils no-code de Google Workspace. "
        "Le champion est autonome : aucune intervention IT n'est necessaire. "
        "Le time-to-value est inferieur a 2 semaines."
    )
    add_pivot_table(
        doc,
        "",
        ["Composant", "Detail"],
        [
            ["Outils", "Gemini (Prompts/Gems), NotebookLM"],
            ["Donnees", "Google Drive / Sheets (statique)"],
            ["Sortie", "Google Docs / Gmail / Google Chat"],
            ["Competence", "Prompt engineering"],
            ["Gouvernance", "Champion seul — aucune intervention IT"],
        ],
    )

    # Medium
    add_sub_heading(doc, "Medium — App & Orchestration")
    add_body_text(
        doc,
        "Les use cases Medium necessitent des outils low-code/semi-code et des appels API. "
        "Le champion travaille avec un support IT local ponctuel, notamment pour les "
        "connexions à des systemes d'entreprise."
    )
    add_pivot_table(
        doc,
        "",
        ["Composant", "Detail"],
        [
            ["Outils", "App Script, AppSheet, AI Studio, Workspace Studio"],
            ["Donnees", "Google Sheets, AppSheet (donnees structurees)"],
            ["Sortie", "AppSheet App, Google Slides/Docs"],
            ["Competence", "Low-code + appels API basiques"],
            ["Gouvernance", "Champion + support IT local ponctuel"],
            ["Alerte IT", "Si connexion SFDC / BI enterprise -> escalade requise"],
        ],
    )

    # Large
    add_sub_heading(doc, "Large — Platform & Agent")
    add_body_text(
        doc,
        "Les use cases Large impliquent des architectures multi-systemes, du code custom, "
        "et souvent des donnees temps reel ou industrielles. Ce sont des projets IT formels "
        "avec budget et gouvernance dedies."
    )
    add_pivot_table(
        doc,
        "",
        ["Composant", "Detail"],
        [
            ["Sources", "SFDC, AVEVA, DCS, SAP, Power BI (Fabric)"],
            ["Ingestion", "Python (Cloud Functions / BigQuery / DataStudio)"],
            ["Moteur IA", "AI Studio / Vertex AI + RAG + Agents Gemini"],
            ["Orchestration", "Advance Coding (backend API / Cloud Run)"],
            ["Frontend", "Web App (AppSheet Pro), Power BI Embedded"],
            ["Competence", "Full-stack + ML engineering"],
            ["Gouvernance", "Projet IT formel + Champion metier + budget dedie"],
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 08 : POINTS D'ATTENTION IT
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "08",
        "Points d'attention IT",
        f"{it_count} use cases necessitant un accompagnement technique dedie",
    )

    add_callout(
        doc,
        "Ces use cases impliquent des dependances systemes (SFDC, SAP, DCS, SCADA, BigQuery, API) "
        "ou ont un profil Large qui depasse les capacites d'un champion seul. "
        "Une coordination avec les equipes IT locales est obligatoire avant tout deploiement.",
        style="warning",
    )

    it_df = df[df["IT_Flag"] != ""].copy()
    it_by_family = (
        it_df.groupby("Family_Label")
        .size()
        .reset_index(name="count")
        .sort_values("count", ascending=False)
    )

    add_pivot_table(
        doc,
        "Points d'attention IT par famille",
        ["Famille fonctionnelle", "Nb use cases IT"],
        [
            [row["Family_Label"], str(row["count"])]
            for _, row in it_by_family.iterrows()
        ],
    )

    add_sub_heading(doc, "Criteres declencheurs d'escalade IT")
    add_bullet(doc, "Connexion a un systeme d'entreprise (ERP, CRM, SCADA, BI enterprise)", "Systeme")
    add_bullet(doc, "Authentification / gestion des droits (SSO, API keys, OAuth)", "Securite")
    add_bullet(doc, "Hebergement hors Google Workspace (Cloud Run, Vertex AI, BDD)", "Infra")
    add_bullet(doc, "Volume de donnees > ce qu'un Google Sheet peut supporter (>100K lignes)", "Donnees")
    add_bullet(doc, "Outil utilise par plus de 10 personnes", "Echelle")

    # Top 15 use cases IT
    p_h = doc.add_paragraph(
        "Selection de use cases a escalader en priorite"
    )
    p_h.style = "Heading 3"
    set_paragraph_spacing(p_h, 12, 4)

    top_it = it_df.nlargest(15, "Score_Total")
    it_rows = []
    for _, row in top_it.iterrows():
        desc = str(row["Use Case Description (Long)"])[:80].strip()
        if len(str(row["Use Case Description (Long)"])) > 80:
            desc += "..."
        it_rows.append(
            [
                str(row["UC_ID"]),
                str(row["Cluster"]),
                str(row["Family_Label"]),
                str(row["Complexity_Tier"]),
                desc,
            ]
        )

    add_pivot_table(
        doc,
        "",
        ["UC_ID", "Cluster", "Famille", "Tier", "Description"],
        it_rows,
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 09 : RECOMMANDATIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "09",
        "Recommandations strategiques",
        "5 axes prioritaires pour industrialiser le portefeuille AI Champions",
    )

    add_body_text(
        doc,
        "Les recommandations suivantes s'appuient sur les constats de l'analyse et visent "
        "à transformer le portefeuille de use cases individuels en une capacité organisationnelle "
        "pérenne. Elles sont classées par priorité d'exécution."
    )

    recommendations = [
        (
            "Industrialiser les Quick Wins (Small)",
            f"{tiers.get('Small', 0)} use cases Small sont deployables immediatement avec peu ou pas de support IT. "
            "Action : constituer une bibliotheque de Gems et de Prompts partages par famille fonctionnelle. "
            "Cible : dupliquer les 10 meilleures pratiques a l'ensemble des clusters en 3 mois. "
            "Indicateur : nombre de Gems partages et taux de reutilisation inter-clusters.",
        ),
        (
            "Structurer l'accompagnement des Medium",
            f"{tiers.get('Medium', 0)} use cases Medium necessitent un encadrement leger "
            "(champion + support IT local). "
            "Action : mettre en place des sprints de 4 semaines avec un champion referent par cluster. "
            "NotebookLM et App Script sont les vecteurs principaux pour ces cas. "
            "Indicateur : taux de passage Medium -> Production en 8 semaines.",
        ),
        (
            "Creer une filiere IT pour les Large",
            f"{tiers.get('Large', 0)} projets strategiques necessitent une equipe projet dediee. "
            "Action : identifier les sponsors metier et les correspondants IT pour chaque projet. "
            "Investissements 3-12 mois avec governance de programme. "
            "Indicateur : nombre de projets Large lances avec sponsor identifie.",
        ),
        (
            "Adresser les points d'attention IT",
            f"{it_count} use cases presentent des dependances systemes critiques. "
            "Action : organiser un atelier de priorisation avec les equipes IT avant tout engagement. "
            "Creer un registre de dependances partage entre les champions et l'IT. "
            "Indicateur : taux de resolution des blocages IT en 30 jours.",
        ),
        (
            "Creer un track Data Champion (F7)",
            "La famille Data Engineering (F7) est la plus dense et la plus complexe. "
            "Les competences Python/BigQuery depassent le profil champion standard. "
            "Action : concevoir un parcours de certification 'Data Champion' avec des modules "
            "specifiques. Indicateur : nombre de Data Champions certifies par cluster.",
        ),
    ]

    for i, (title, body) in enumerate(recommendations, start=1):
        p_rec = doc.add_paragraph()
        badge_r = p_rec.add_run(f" 0{i} ")
        badge_r.font.name = "Poppins"
        badge_r.font.bold = True
        badge_r.font.size = Pt(12)
        badge_r.font.color.rgb = PYL_NAVY_DARK
        highlight_run(badge_r)
        p_rec.add_run("  ")
        t_run = p_rec.add_run(title)
        t_run.font.name = "Poppins"
        t_run.font.bold = True
        t_run.font.size = Pt(14)
        t_run.font.color.rgb = PYL_NAVY_DARK
        set_paragraph_spacing(p_rec, 14, 4)

        p_body = doc.add_paragraph(body)
        p_body.style = "Normal"
        set_paragraph_spacing(p_body, 0, 10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 10 : BONNES PRATIQUES
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "10",
        "Guide des bonnes pratiques",
        "Recommandations methodologiques pour les AI Champions",
    )

    add_body_text(
        doc,
        "Les champions sont des bidouilleurs metiers, pas des developpeurs. "
        "Sans cadre minimal, leurs realisations restent fragiles, non maintenables "
        "et impossibles a transmettre. Les recommandations ci-dessous visent a instaurer "
        "une discipline legere mais structurante."
    )

    # 10.1 — Gestion des prompts
    add_sub_heading(doc, "10.1 — Gestion et versioning des prompts")
    add_body_text(
        doc,
        "Les prompts Gemini sont des actifs metier au meme titre qu'une procedure operationnelle. "
        "Ils doivent etre stockes, versionnes et documentes."
    )
    add_pivot_table(
        doc,
        "",
        ["Pratique", "Comment faire (sans IT)"],
        [
            ["Stocker les prompts", "Google Doc dedie par use case, pas dans le chat Gemini"],
            ["Versionner", "Naming : Prompt_V1.0_AAAA-MM-JJ dans le titre du Doc"],
            ["Documenter le contexte", "En-tete : objectif, audience, exemples input/output"],
            ["Centraliser", "Google Drive partage AI Champions / Prompts Library"],
            ["Tester avant de deployer", "3 exemples reels avant de partager un Gem"],
        ],
    )

    # 10.2 — App Script
    add_sub_heading(doc, "10.2 — Bonnes pratiques App Script")
    add_body_text(
        doc,
        "App Script est l'outil le plus utilise du portefeuille. "
        "Sans discipline minimale, les scripts deviennent impossibles a maintenir."
    )
    add_pivot_table(
        doc,
        "",
        ["Regle", "Pourquoi"],
        [
            ["1 script = 1 fichier nomme", "Pas de scripts eparpilles dans des Sheets au hasard"],
            ["Commenter les blocs", "Le champion suivant doit comprendre"],
            ["Pas de secrets en clair", "Utiliser PropertiesService.getScriptProperties()"],
            ["Creer un README", "Expliquer ce que fait le script, comment le lancer"],
            ["Sauvegarder avant modification", "Copier dans un Doc Archive avec la date"],
            ["Tester sur donnees fictives", "Jamais directement sur la production"],
        ],
    )

    # 10.3 — NotebookLM
    add_sub_heading(doc, "10.3 — Bonnes pratiques NotebookLM")
    add_pivot_table(
        doc,
        "",
        ["Regle", "Pourquoi"],
        [
            ["Documenter les sources et leur date", "NotebookLM n'est pas a jour automatiquement"],
            ["Sources structurees (PDF, Doc)", "Ameliore la qualite des reponses"],
            ["1 Notebook par perimetre", "Eviter les notebooks fourre-tout"],
            ["Rafraichir les sources trimestriellement", "Sources obsoletes = reponses incorrectes"],
        ],
    )

    # 10.4 — Kit Champion
    add_sub_heading(doc, "10.4 — Gouvernance legere : le Kit Champion")
    add_body_text(
        doc,
        "Pour chaque use case cree, le champion devrait maintenir un Kit Champion minimal, "
        "stocke dans le Google Drive partage de l'equipe AI Champions du cluster :"
    )
    add_bullet(doc, "Description, audience, how-to, contacts", "README.md")
    add_bullet(doc, "Historique des prompts versionnes", "Prompts_Vx.x.md")
    add_bullet(doc, "Copie du App Script (si applicable)", "Script_Vx.x.gs")
    add_bullet(doc, "Exemples d'inputs / outputs valides", "Tests.md")
    add_bullet(doc, "Historique des modifications", "Changelog.md")

    # 10.5 — Quand escalader
    add_sub_heading(doc, "10.5 — Quand escalader vers l'IT ?")
    add_pivot_table(
        doc,
        "",
        ["Signal", "Action recommandee"],
        [
            ["Besoin d'acceder a un systeme enterprise", "Contacter l'IT pour une API / connexion securisee"],
            ["Donnees sensibles (RH, financieres, clients)", "Verifier avec le DPO / Legal"],
            ["Le script plante regulierement", "Qualifier en Large, escalader"],
            ["Utilise par plus de 10 personnes", "Passer en Medium/Large, documenter formellement"],
            ["Besoin de fiabilite 24/7", "Hors perimetre champion -> projet IT"],
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    #  SECTION 11 : FEUILLE DE ROUTE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_header(
        doc,
        "11",
        "Feuille de route",
        "Jalons de mise en oeuvre sur 12 mois",
    )

    add_body_text(
        doc,
        "La feuille de route ci-dessous propose un cadencement progressif, "
        "en commencant par les Quick Wins a fort impact et faible effort, "
        "puis en structurant l'accompagnement des use cases Medium, "
        "avant de lancer les projets strategiques Large."
    )

    add_pivot_table(
        doc,
        "Roadmap d'industrialisation",
        ["Jalon", "Actions", "Indicateur de succes"],
        [
            [
                "M+1",
                "Atelier de priorisation Quick Wins par cluster. "
                "Identification des 10 Gems partageables.",
                f"10 Quick Wins choisis sur {tiers.get('Small', 0)} disponibles",
            ],
            [
                "M+3",
                "Bibliotheque Gems V1 disponible. "
                "Premiers sprints Medium lances (2 par cluster).",
                "Bibliotheque accessible. 6 sprints Medium lances.",
            ],
            [
                "M+6",
                "50% des Medium en production. "
                "Track Data Champion defini. Registre IT operationnel.",
                f"{tiers.get('Medium', 0) // 2} Medium en production",
            ],
            [
                "M+12",
                "Projets Large lances avec sponsors IT identifies. "
                "Certification Data Champion operationnelle.",
                f"{tiers.get('Large', 0)} projets Large avec budget. "
                "10+ Data Champions certifies.",
            ],
        ],
    )

    add_callout(
        doc,
        "Cette feuille de route est indicative et devra etre adaptee en fonction "
        "des priorites de chaque cluster et des ressources IT disponibles. "
        "Un point d'avancement trimestriel est recommande.",
        style="info",
    )

    # Footer final
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run(
        f"(c) Copyright {YEAR} Pyl.Tech  |  Confidentiel  |  Air Liquide AI Champions"
    )
    r_footer.font.name = "Poppins"
    r_footer.font.size = Pt(9)
    r_footer.font.color.rgb = PYL_TEAL_BLUE
    set_paragraph_spacing(p_footer, 24, 0)


# ── Point d'entree ────────────────────────────────────────────────────────────


def main() -> None:
    print(f"Lecture du catalogue : {CATALOG_FILE}")  # noqa
    if not CATALOG_FILE.exists():
        print(f"Catalogue introuvable : {CATALOG_FILE}")  # noqa
        print("   Lancez d'abord : python3 src/generate_catalog.py")  # noqa
        sys.exit(1)

    df = pd.read_excel(CATALOG_FILE, sheet_name="Catalogue", header=0)
    print(f"   {len(df)} use cases charges")  # noqa

    doc = Document()
    apply_base_styles(doc)
    add_cover_page(doc)
    build_document(doc, df)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\nRapport Word genere : {OUTPUT_FILE}")  # noqa
    print(f"   Taille : {OUTPUT_FILE.stat().st_size // 1024} KB")  # noqa


if __name__ == "__main__":
    main()
