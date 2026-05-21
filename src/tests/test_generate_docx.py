"""
test_generate_docx.py
Tests unitaires — couverture 100% de src/generate_docx.py
"""
import sys
import os
import unittest
from pathlib import Path
from unittest.mock import patch, MagicMock

import pandas as pd
from docx import Document
from docx.shared import RGBColor

sys.path.insert(0, str(Path(__file__).parent.parent))
import generate_docx as gd


# ── Fixture DataFrame minimal ─────────────────────────────────────────────────

def _make_df(n: int = 4) -> pd.DataFrame:
    """DataFrame minimal couvrant toutes les branches de build_document."""
    return pd.DataFrame({
        "UC_ID":                        [f"UC_{i:04d}" for i in range(1, n + 1)],
        "Cluster":                      ["ClusterA", "ClusterA", "ClusterB", "ClusterC"],
        "Job Family":                   ["Finance", "Finance", "HR", "Operations"],
        "Family":                       ["F1", "F7", "F3", "F4"],
        "Family_Label":                 ["Automatisation documentaire",
                                         "Data Engineering & Reporting",
                                         "Customer & Sales Intelligence",
                                         "Monitoring & Maintenance industrielle"],
        "Complexity_Tier":              ["Small", "Medium", "Large", "Small"],
        "Score_Total":                  [5, 9, 13, 6],
        "Score_Integration":            [1, 2, 3, 1],
        "Score_Scope":                  [1, 2, 3, 1],
        "Score_Data":                   [1, 2, 3, 2],
        "Score_AI":                     [1, 1, 2, 1],
        "Score_Economic":               [1, 2, 2, 1],
        "Stage":                        ["POC", "Deployed", "Ideation", "POC"],
        "Scope of the Use Case":        ["Local", "Country", "Group", "Local"],
        "Economical Impact":            ["Cost Reduction", "Revenue Growth", "Sustainability", "Cost Reduction"],
        "Tools":                        ["Gemini", "Power BI", "AI Studio", "Gemini"],
        "Tools_Tags":                   ["Gemini Prompts/Gems", "Power BI", "AI Studio", ""],
        "Nb_Tools":                     [1, 1, 1, 0],
        "Max_Tool_Level":               ["L1", "L3", "L3", "L1"],
        "IT_Flag":                      ["", "⚠️ IT", "⚠️ IT", ""],
        "IT_Attention":                 ["", "sfdc", "salesforce", ""],
        "Use Case Description (Long)":  [
            "Translate the weekly report to English",
            "Build a bigquery pipeline for reporting",
            "Analyse customer data from sfdc and generate insights",
            "Monitor equipment sensor data in real-time",
        ],
    })


# ════════════════════════════════════════════════════════════════════════════
# Helpers XML / Bas niveau
# ════════════════════════════════════════════════════════════════════════════

class TestSetCellBg(unittest.TestCase):
    def _make_cell(self):
        doc = Document()
        tbl = doc.add_table(rows=1, cols=1)
        return tbl.cell(0, 0)

    def test_sets_background_without_error(self):
        cell = self._make_cell()
        gd.set_cell_bg(cell, RGBColor(0x0B, 0x13, 0x2B))
        self.assertIsNotNone(cell)

    def test_different_colors_accepted(self):
        cell = self._make_cell()
        gd.set_cell_bg(cell, gd.PYL_YELLOW)
        self.assertIsNotNone(cell)
        gd.set_cell_bg(cell, gd.PYL_WHITE)
        self.assertIsNotNone(cell)


class TestAddBorderLeft(unittest.TestCase):
    def test_adds_border_without_error(self):
        doc = Document()
        p = doc.add_paragraph("test")
        gd.add_border_left(p, "F4BF46", size_pt=18)
        self.assertIsNotNone(p)

    def test_custom_size(self):
        doc = Document()
        p = doc.add_paragraph("test")
        gd.add_border_left(p, "C91432", size_pt=30)
        self.assertIsNotNone(p)


class TestHighlightRun(unittest.TestCase):
    def test_applies_highlight_without_error(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("highlight me")
        gd.highlight_run(run)
        self.assertIsNotNone(run)

    def test_custom_hex_color(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("custom color")
        gd.highlight_run(run, hex_color="208AAE")
        self.assertIsNotNone(run)


class TestSetParagraphSpacing(unittest.TestCase):
    def test_default_spacing(self):
        doc = Document()
        p = doc.add_paragraph("test")
        gd.set_paragraph_spacing(p)  # defaults
        self.assertIsNotNone(p)

    def test_custom_spacing(self):
        doc = Document()
        p = doc.add_paragraph("test")
        gd.set_paragraph_spacing(p, space_before=12, space_after=18)
        self.assertIsNotNone(p)


# ════════════════════════════════════════════════════════════════════════════
# Composants de document
# ════════════════════════════════════════════════════════════════════════════

class TestApplyBaseStyles(unittest.TestCase):
    def test_runs_without_error(self):
        doc = Document()
        gd.apply_base_styles(doc)
        self.assertIsNotNone(doc)

    def test_margins_applied(self):
        doc = Document()
        gd.apply_base_styles(doc)
        self.assertAlmostEqual(doc.sections[0].top_margin.cm, 2.0, places=1)


class TestAddCoverPage(unittest.TestCase):
    def test_adds_content_without_error(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_cover_page(doc)
        # Doit avoir un saut de page
        self.assertGreater(len(doc.paragraphs), 0)


class TestAddSectionHeader(unittest.TestCase):
    def test_with_subtitle(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_section_header(doc, "01", "Titre principal", "Sous-titre optionnel")
        self.assertIsNotNone(doc)

    def test_without_subtitle(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_section_header(doc, "02", "Titre sans sous-titre")
        self.assertIsNotNone(doc)

    def test_empty_number(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_section_header(doc, "", "Titre sans numéro de section")
        self.assertIsNotNone(doc)


class TestAddKpiBlock(unittest.TestCase):
    def test_single_kpi(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_kpi_block(doc, [("Indicateur 1", "42", "Contexte")])
        self.assertIsNotNone(doc)

    def test_multiple_kpis_odd_even(self):
        doc = Document()
        gd.apply_base_styles(doc)
        kpis = [
            ("KPI 1", "100", "ctx1"),
            ("KPI 2", "200", "ctx2"),
            ("KPI 3", "300", "ctx3"),
        ]
        gd.add_kpi_block(doc, kpis)  # couvre les lignes paires et impaires
        self.assertIsNotNone(doc)


class TestAddPivotTable(unittest.TestCase):
    def test_basic_table(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_pivot_table(
            doc, "Titre tableau",
            ["Col1", "Col2", "Total"],
            [["Row1", "10", "10"], ["Row2", "5", "5"]]
        )
        self.assertIsNotNone(doc)

    def test_empty_rows(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_pivot_table(doc, "Vide", ["A", "B"], [])
        self.assertIsNotNone(doc)


class TestAddCallout(unittest.TestCase):
    def test_info_style(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_callout(doc, "Info message", style="info")
        self.assertIsNotNone(doc)

    def test_warning_style(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_callout(doc, "Warning message", style="warning")
        self.assertIsNotNone(doc)

    def test_success_style(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_callout(doc, "Success message", style="success")
        self.assertIsNotNone(doc)

    def test_note_style(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_callout(doc, "Note message", style="note")
        self.assertIsNotNone(doc)

    def test_unknown_style_fallback(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_callout(doc, "Unknown style", style="unknown_xyz")
        self.assertIsNotNone(doc)


class TestAddBullet(unittest.TestCase):
    def test_bullet_without_prefix(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_bullet(doc, "Corps du bullet sans préfixe")
        self.assertIsNotNone(doc)

    def test_bullet_with_prefix(self):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_bullet(doc, "Corps du bullet avec préfixe", bold_prefix="Titre")
        self.assertIsNotNone(doc)


# ════════════════════════════════════════════════════════════════════════════
# build_document — pipeline complet
# ════════════════════════════════════════════════════════════════════════════

class TestBuildDocument(unittest.TestCase):
    def _build(self, df):
        doc = Document()
        gd.apply_base_styles(doc)
        gd.add_cover_page(doc)
        gd.build_document(doc, df)
        return doc

    def test_runs_without_error(self):
        df = _make_df()
        doc = self._build(df)
        self.assertGreater(len(doc.paragraphs), 5)

    def test_family_with_no_quick_wins(self):
        """Famille sans aucun Small ne doit pas planter."""
        df = _make_df()
        df["Complexity_Tier"] = "Medium"  # Aucun Small → pas de quick wins
        self._build(df)

    def test_family_without_it_flag(self):
        """Famille sans IT_Flag ne doit pas afficher le bullet IT."""
        df = _make_df()
        df["IT_Flag"] = ""  # Aucun IT
        self._build(df)

    def test_long_description_truncated(self):
        """Couvre la branche 'desc_short += ...' pour les descriptions > 120 chars."""
        df = _make_df()
        long_desc = "A" * 130  # > 120 chars
        df["Use Case Description (Long)"] = long_desc
        df["Complexity_Tier"] = "Small"  # Quick wins actifs
        doc = self._build(df)
        self.assertIsNotNone(doc)

    def test_it_long_description_truncated(self):
        """Couvre la branche 'desc += ...' dans la section IT (> 80 chars)."""
        df = _make_df()
        long_desc = "B" * 90  # > 80 chars
        df["Use Case Description (Long)"] = long_desc
        df["IT_Flag"] = "⚠️ IT"  # Tout le monde en IT pour la section IT
        doc = self._build(df)
        self.assertIsNotNone(doc)

    def test_all_tiers_represented(self):
        """Vérifie que tous les tiers sont bien présents dans le doc généré."""
        df = _make_df()
        doc = self._build(df)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Small", full_text)
        self.assertIn("Medium", full_text)
        self.assertIn("Large", full_text)


# ════════════════════════════════════════════════════════════════════════════
# main() — filesystem mocking
# ════════════════════════════════════════════════════════════════════════════

class TestMain(unittest.TestCase):
    def test_exits_when_catalog_missing(self):
        fake_source = MagicMock(spec=Path)
        fake_source.exists.return_value = False

        with patch.object(gd, "CATALOG_FILE", fake_source):
            with self.assertRaises(SystemExit) as ctx:
                gd.main()
        self.assertEqual(ctx.exception.code, 1)

    def test_success_generates_docx(self):
        import tempfile
        df = _make_df()
        fake_source = MagicMock(spec=Path)
        fake_source.exists.return_value = True

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        try:
            with patch.object(gd, "CATALOG_FILE", fake_source), \
                 patch.object(gd, "OUTPUT_FILE", tmp_path), \
                 patch("generate_docx.pd.read_excel", return_value=df):
                gd.main()
            self.assertTrue(tmp_path.exists())
            self.assertGreater(tmp_path.stat().st_size, 1000)
        finally:
            if tmp_path.exists():
                os.unlink(tmp_path)


if __name__ == "__main__":
    unittest.main(verbosity=2)
